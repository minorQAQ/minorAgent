"""文档操作工具：对 plain_text / tabular / docx / pptx / pdf 文件进行增删改查。

plain_text:  行级粒度（与代码 agent 一致）
tabular:     单元格级粒度（CSV / XLSX）
docx:        段落级粒度（文本段落）；表格操作粒度到单元格
pptx:        只读 + 删除（按幻灯片提取文本+内嵌图片+表格）
pdf:         只读 + 删除
"""

from __future__ import annotations

import csv
import json
import os
import pathlib
import re
from typing import Literal, Optional

from langchain.tools import BaseTool
from pydantic import BaseModel, Field

from agent.utils.agent_utils import DOC_EXT_MAP, Documents_process
from agent.utils.image_utils import IMAGE_FILE_EXTENSIONS
from agent.utils.agent_utils import AUDIO_FILE_EXTENSIONS, audio_file_to_text


def _classify_by_ext(file_path: str) -> str:
    """根据扩展名推断文档类型（文件可不存在，用于 create 操作）。"""
    ext = pathlib.Path(file_path).suffix.lower()
    mapped = DOC_EXT_MAP.get(ext, "")
    if mapped:
        return mapped
    if ext in IMAGE_FILE_EXTENSIONS:
        return "image"
    if ext in AUDIO_FILE_EXTENSIONS:
        return "audio"
    return ""


def _classify(file_path: str) -> str:
    """获取文档类型：文件存在时用 Documents_process 分类，否则按扩展名推断。"""
    if os.path.isfile(file_path):
        return Documents_process.doc_classify(file_path)
    return _classify_by_ext(file_path)


class DocToolInput(BaseModel):
    file_path: str = Field(..., description="文档的完整路径")
    action: Literal["read", "create", "update", "delete"] = Field(..., description="操作类型")

    # ---- create / update 通用 ----
    content: Optional[str] = Field(
        None,
        description=(
            "create 时写入的文本内容。"
            "plain_text/docx: 直接写入文本。"
            "tabular(.csv): CSV 格式文本。"
            "tabular(.xlsx): JSON 二维数组字符串，如 '[[\"A\",\"B\"],[\"1\",\"2\"]]'。"
            "tabular(.xls): 仅支持 read/delete，不支持 create/update。建议转 .xlsx。"
        ),
    )

    # ---- plain_text 行级操作 ----
    line_number: Optional[int] = Field(None, description="plain_text: 目标行号(1-indexed)")
    new_line_text: Optional[str] = Field(None, description="plain_text: 替换该行时的新文本")
    insert_before_line: Optional[int] = Field(None, description="plain_text/docx: 在此行/段落索引前插入(1-indexed)")
    insert_text: Optional[str] = Field(None, description="plain_text/docx: 插入的新文本内容")

    # ---- tabular 单元格操作 ----
    sheet_name: Optional[str] = Field(None, description="tabular(.xlsx): sheet 名称，默认第一个 sheet")
    row: Optional[int] = Field(None, description="tabular: 行号(1-indexed，1=表头)")
    col: Optional[int] = Field(None, description="tabular: 列号(1-indexed)")
    cell_value: Optional[str] = Field(None, description="tabular: 单元格新值；为空则清空该单元格")

    # ---- docx 段落/表格操作 ----
    paragraph_index: Optional[int] = Field(None, description="docx: 段落索引(0-indexed)")
    new_paragraph_text: Optional[str] = Field(None, description="docx: 替换段落时的新文本")
    insert_paragraph_at: Optional[int] = Field(None, description="docx: 在此索引前插入新段落(0-indexed)")
    table_index: Optional[int] = Field(None, description="docx: 表格索引(0-indexed)")
    docx_table_row: Optional[int] = Field(None, description="docx: 表格行号(0-indexed)")
    docx_table_col: Optional[int] = Field(None, description="docx: 表格列号(0-indexed)")
    docx_table_cell_value: Optional[str] = Field(None, description="docx: 表格单元格新值")

    # ---- docx 图片操作 ----
    image_paths: Optional[list[str]] = Field(
        None,
        description="docx create: 图片文件路径列表，按顺序插入到文本段落后。每个路径必须指向存在的图片文件(png/jpg等)。",
    )
    insert_image_path: Optional[str] = Field(
        None,
        description="docx update: 要插入的图片文件路径。需配合 insert_paragraph_at 指定插入位置。",
    )

    # ---- read 内容截取控制 ----
    max_chars: Optional[int] = Field(None, description="read 时最多返回的字符数。None 表示全部返回")
    max_images: Optional[int] = Field(None, description="read 时 Documents_process 最多返回的图片数。None 表示全部返回")
    from_end: bool = Field(False, description="read 截取时是否从末尾开始。False 从头，True 从尾")

    # ---- plain_text grep ----
    grep_pattern: Optional[str] = Field(None, description="plain_text read 时：正则过滤行，仅返回匹配的行(含行号)")


def _apply_text_truncation(text: str, max_chars: int | None, from_end: bool) -> str:
    """对文本按 max_chars 和 from_end 截取，返回截取后文本 + 截断提示。"""
    if max_chars is None or max_chars <= 0 or len(text) <= max_chars:
        return text
    if from_end:
        truncated = text[-max_chars:]
        return f"...(截取了末尾 {max_chars} 字符，原文共 {len(text)} 字符)\n{truncated}"
    else:
        truncated = text[:max_chars]
        return f"{truncated}\n...(截取了前 {max_chars} 字符，原文共 {len(text)} 字符)"


def _read_plain_text(file_path: str, grep_pattern: str | None = None,
                     max_chars: int | None = None, from_end: bool = False) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    total_lines = len(lines)

    if grep_pattern:
        try:
            pat = re.compile(grep_pattern)
        except re.error as e:
            return f"[ERROR] 无效的正则表达式: {e}"
        matched = [(i, line) for i, line in enumerate(lines) if pat.search(line)]
        if not matched:
            return f"=== {file_path} (grep: '{grep_pattern}') - 无匹配 ===\n"
        numbered = "".join(f"{i + 1:>4}| {line}" for i, line in matched)
        result = f"=== {file_path} (grep: '{grep_pattern}', {len(matched)}/{total_lines} lines) ===\n{numbered}"
        return _apply_text_truncation(result, max_chars, from_end)

    numbered = "".join(f"{i + 1:>4}| {line}" for i, line in enumerate(lines))
    result = f"=== {file_path} ({total_lines} lines) ===\n{numbered}"
    return _apply_text_truncation(result, max_chars, from_end)


def _create_plain_text(file_path: str, content: str) -> str:
    if os.path.exists(file_path):
        return f"[ERROR] 文件已存在: {file_path}"
    # 自动创建缺失的父目录（配合 plan 模式写入 .minorAgent 等场景）
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return f"[OK] 已创建文件: {file_path}"


def _update_plain_text(file_path: str, line_number: int | None, new_line_text: str | None,
                       insert_before_line: int | None, insert_text: str | None) -> str:
    if not os.path.isfile(file_path):
        return f"[ERROR] 文件不存在: {file_path}"
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    if insert_before_line is not None and insert_text is not None:
        idx = insert_before_line - 1
        if idx < 0:
            idx = 0
        if idx > len(lines):
            idx = len(lines)
        lines.insert(idx, insert_text.rstrip("\n") + "\n")
        with open(file_path, "w", encoding="utf-8") as f:
            f.writelines(lines)
        return f"[OK] 已在第 {insert_before_line} 行前插入新行: {file_path}"

    if line_number is not None and new_line_text is not None:
        idx = line_number - 1
        if idx < 0 or idx >= len(lines):
            return f"[ERROR] 行号 {line_number} 超出范围 (1-{len(lines)})"
        lines[idx] = new_line_text.rstrip("\n") + "\n"
        with open(file_path, "w", encoding="utf-8") as f:
            f.writelines(lines)
        return f"[OK] 已替换第 {line_number} 行: {file_path}"

    if line_number is not None and new_line_text is None and insert_before_line is None:
        # delete line
        idx = line_number - 1
        if idx < 0 or idx >= len(lines):
            return f"[ERROR] 行号 {line_number} 超出范围 (1-{len(lines)})"
        removed = lines.pop(idx).rstrip("\n")
        with open(file_path, "w", encoding="utf-8") as f:
            f.writelines(lines)
        return f"[OK] 已删除第 {line_number} 行: '{removed[:80]}'"

    return "[ERROR] update 操作缺少必要参数 (line_number + new_line_text 或 insert_before_line + insert_text)"


def _delete_file(file_path: str) -> str:
    if not os.path.isfile(file_path):
        return f"[ERROR] 文件不存在: {file_path}"
    os.remove(file_path)
    return f"[OK] 已删除文件: {file_path}"


# ---------- tabular ----------
def _read_tabular(file_path: str, sheet_name: str | None = None,
                  max_chars: int | None = None, from_end: bool = False) -> str:
    ext = pathlib.Path(file_path).suffix.lower()
    if ext == ".csv":
        with open(file_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
        if not rows:
            return f"=== {file_path} (empty) ==="
        header = " | ".join(rows[0])
        sep = "-" * len(header)
        body = "\n".join(f"{i + 1:>4}| {' | '.join(row)}" for i, row in enumerate(rows))
        result = f"=== {file_path} (CSV, {len(rows)} rows) ===\nRow | {header}\n{sep}\n{body}"
        return _apply_text_truncation(result, max_chars, from_end)

    elif ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        names = wb.sheetnames
        target = sheet_name or (names[0] if names else None)
        if target is None or target not in names:
            wb.close()
            return f"[ERROR] sheet '{sheet_name}' 不存在，可用: {names}"
        ws = wb[target]
        rows = [[str(c) if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
        wb.close()
        if not rows:
            return f"=== {file_path} [{target}] (empty) ==="
        header = " | ".join(rows[0])
        sep = "-" * len(header)
        body = "\n".join(f"{i + 1:>4}| {' | '.join(row)}" for i, row in enumerate(rows))
        result = f"=== {file_path} [{target}] ({len(rows)} rows) ===\nRow | {header}\n{sep}\n{body}"
        return _apply_text_truncation(result, max_chars, from_end)

    elif ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(file_path)
        target = sheet_name or (wb.sheet_names()[0] if wb.sheet_names() else None)
        if target is None or target not in wb.sheet_names():
            return f"[ERROR] sheet '{sheet_name}' 不存在，可用: {list(wb.sheet_names())}"
        ws = wb.sheet_by_name(target)
        rows = []
        for r in range(ws.nrows):
            row_vals = []
            for c in range(ws.ncols):
                v = ws.cell_value(r, c)
                if ws.cell_type(r, c) == xlrd.XL_CELL_DATE:
                    try:
                        v = xlrd.xldate_as_datetime(v, wb.datemode)
                    except Exception:
                        pass
                row_vals.append(str(v) if v is not None else "")
            rows.append(row_vals)
        if not rows:
            return f"=== {file_path} [{target}] (empty) ==="
        header = " | ".join(rows[0])
        sep = "-" * len(header)
        body = "\n".join(f"{i + 1:>4}| {' | '.join(row)}" for i, row in enumerate(rows))
        result = f"=== {file_path} [{target}] ({len(rows)} rows) ===\nRow | {header}\n{sep}\n{body}"
        return _apply_text_truncation(result, max_chars, from_end)

    return f"[ERROR] 不支持的 tabular 格式: {ext}"


def _create_tabular(file_path: str, content: str) -> str:
    if os.path.exists(file_path):
        return f"[ERROR] 文件已存在: {file_path}"
    # 自动创建缺失的父目录
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    ext = pathlib.Path(file_path).suffix.lower()
    if ext == ".csv":
        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            f.write(content)
        return f"[OK] 已创建 CSV 文件: {file_path}"
    elif ext == ".xlsx":
        from openpyxl import Workbook
        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            return f"[ERROR] content 需为 JSON 二维数组字符串，如 '[[\"A\",\"B\"],[\"1\",\"2\"]]'"
        if not isinstance(data, list) or not all(isinstance(r, list) for r in data):
            return f"[ERROR] content 需为 JSON 二维数组字符串"
        wb = Workbook()
        ws = wb.active
        for row in data:
            ws.append(row)
        wb.save(file_path)
        return f"[OK] 已创建 XLSX 文件: {file_path}"

    elif ext == ".xls":
        return "[ERROR] 不支持创建 .xls 文件（xlrd 为只读库）。请使用 .xlsx 扩展名创建。"
    return f"[ERROR] 不支持的 tabular 格式: {ext}"


def _update_tabular(file_path: str, row: int | None, col: int | None,
                    cell_value: str | None, sheet_name: str | None) -> str:
    if not os.path.isfile(file_path):
        return f"[ERROR] 文件不存在: {file_path}"
    ext = pathlib.Path(file_path).suffix.lower()

    if row is None or col is None:
        return "[ERROR] tabular update 需要 row 和 col 参数"

    if ext == ".csv":
        with open(file_path, "r", encoding="utf-8-sig", newline="") as f:
            rows = list(csv.reader(f))
        r_idx = row - 1
        c_idx = col - 1
        if r_idx < 0 or r_idx >= len(rows):
            return f"[ERROR] 行号 {row} 超出范围 (1-{len(rows)})"
        if c_idx < 0 or c_idx >= len(rows[r_idx]):
            # 自动扩展列
            while len(rows[r_idx]) <= c_idx:
                rows[r_idx].append("")
        rows[r_idx][c_idx] = cell_value if cell_value is not None else ""
        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(rows)
        return f"[OK] 已更新 CSV 单元格 ({row},{col}): {file_path}"

    elif ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        target = sheet_name or wb.sheetnames[0]
        if target not in wb.sheetnames:
            wb.close()
            return f"[ERROR] sheet '{target}' 不存在"
        ws = wb[target]
        ws.cell(row=row, column=col, value=cell_value if cell_value is not None else "")
        wb.save(file_path)
        wb.close()
        return f"[OK] 已更新 XLSX 单元格 ({row},{col}) [{target}]: {file_path}"

    elif ext == ".xls":
        return "[ERROR] 不支持修改 .xls 文件（xlrd 为只读库）。请先将文件转为 .xlsx 格式再操作。"

    return f"[ERROR] 不支持的 tabular 格式: {ext}"


# ---------- docx ----------
def _read_docx(file_path: str, max_chars: int | None = None, from_end: bool = False,
               max_images: int | None = None) -> tuple[str, list[bytes]]:
    text, images = Documents_process.process_docx(file_path)
    # 截取图片
    if max_images is not None and max_images >= 0 and len(images) > max_images:
        images = images[:max_images] if not from_end else images[-max_images:]
    img_note = f"\n(含 {len(images)} 张内嵌图片)" if images else ""
    result = f"=== {file_path} ===\n{text}{img_note}"
    result = _apply_text_truncation(result, max_chars, from_end)
    return result, images


def _strip_markdown(text: str) -> str:
    """清洗 Markdown 标记，返回纯文本段落列表。"""
    lines: list[str] = []
    for line in text.split("\n"):
        # 跳过图片语法 ![alt](url)
        if re.match(r"^!\[.*\]\(.*\)$", line.strip()):
            continue
        # 去除链接 [text](url) → text
        line = re.sub(r"\[([^\]]*)\]\([^\)]*\)", r"\1", line)
        # 去除粗体/斜体 **text** / *text*
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"\*(.+?)\*", r"\1", line)
        # 去除行内代码 `code`
        line = re.sub(r"`([^`]+)`", r"\1", line)
        # 去除水平分割线
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            continue
        # 去除标题标记
        line = re.sub(r"^#{1,6}\s+", "", line)
        # 去除无序列表标记
        line = re.sub(r"^[\s]*[-*+]\s+", "", line)
        # 去除有序列表标记
        line = re.sub(r"^\s*\d+\.\s+", "", line)
        lines.append(line)
    return "\n".join(lines)


def _create_docx(file_path: str, content: str, image_paths: list[str] | None = None) -> str:
    """创建 DOCX 文件，可附带插入图片。

    输入:
        file_path: 目标文件路径。
        content: 文本内容（可为 Markdown 格式，自动清洗标记后写入）。
        image_paths: 可选，图片文件路径列表，按顺序插入到所有文本段落之后。

    输出:
        操作结果描述。
    """
    if os.path.exists(file_path):
        return f"[ERROR] 文件已存在: {file_path}"
    from docx import Document
    from docx.shared import Inches
    # 自动创建缺失的父目录
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc = Document()
    # 清洗 Markdown 标记后写入文本段落
    cleaned = _strip_markdown(content)
    for para_text in cleaned.split("\n"):
        stripped = para_text.strip()
        if not stripped:
            continue  # 跳过空行
        doc.add_paragraph(stripped)
    # 插入图片
    img_count = 0
    if image_paths:
        for img_path in image_paths:
            if not os.path.isfile(img_path):
                continue
            try:
                doc.add_picture(img_path, width=Inches(5.5))
                # 为图片添加一个空段落作为间距
                last_para = doc.add_paragraph("")
                img_count += 1
            except Exception:
                continue
    doc.save(file_path)
    extra = f"，并插入 {img_count} 张图片" if img_count > 0 else ""
    return f"[OK] 已创建 DOCX 文件: {file_path}{extra}"


def _update_docx(file_path: str, paragraph_index: int | None, new_paragraph_text: str | None,
                 insert_paragraph_at: int | None, insert_text: str | None,
                 table_index: int | None, docx_table_row: int | None,
                 docx_table_col: int | None, docx_table_cell_value: str | None,
                 insert_image_path: str | None = None) -> str:
    if not os.path.isfile(file_path):
        return f"[ERROR] 文件不存在: {file_path}"
    from docx import Document

    doc = Document(file_path)

    # ---- 表格单元格操作 ----
    if table_index is not None and docx_table_row is not None and docx_table_col is not None:
        if table_index < 0 or table_index >= len(doc.tables):
            return f"[ERROR] 表格索引 {table_index} 超出范围 (0-{len(doc.tables) - 1})"
        table = doc.tables[table_index]
        if docx_table_row < 0 or docx_table_row >= len(table.rows):
            return f"[ERROR] 表格行 {docx_table_row} 超出范围 (0-{len(table.rows) - 1})"
        row = table.rows[docx_table_row]
        if docx_table_col < 0 or docx_table_col >= len(row.cells):
            return f"[ERROR] 表格列 {docx_table_col} 超出范围 (0-{len(row.cells) - 1})"
        cell = row.cells[docx_table_col]
        cell.text = docx_table_cell_value if docx_table_cell_value is not None else ""
        doc.save(file_path)
        return f"[OK] 已更新 DOCX 表格 ({table_index}, {docx_table_row}, {docx_table_col}): {file_path}"

    # ---- 段落列表（供插入图片与后续段落操作共用）----
    paragraphs = doc.paragraphs
    para_count = len(paragraphs)

    # ---- 插入图片 ----
    if insert_paragraph_at is not None and insert_image_path is not None:
        if not os.path.isfile(insert_image_path):
            return f"[ERROR] 图片文件不存在: {insert_image_path}"
        if insert_paragraph_at < 0:
            insert_paragraph_at = 0
        if insert_paragraph_at > para_count:
            insert_paragraph_at = para_count
        from docx.shared import Inches
        try:
            if insert_paragraph_at == para_count:
                # 末尾追加
                doc.add_picture(insert_image_path, width=Inches(5.5))
                doc.add_paragraph("")
            else:
                ref_para = paragraphs[insert_paragraph_at]
                # 在参考段落前插入一个新的空段落，然后在空段落中插入图片
                new_para = ref_para.insert_paragraph_before("")
                run = new_para.add_run()
                run.add_picture(insert_image_path, width=Inches(5.5))
            doc.save(file_path)
            return f"[OK] 已在 DOCX 段落 {insert_paragraph_at} 前插入图片: {file_path}"
        except Exception as e:
            return f"[ERROR] 插入图片失败: {str(e)}"

    # ---- 段落操作 ----
    # 删除段落
    if paragraph_index is not None and new_paragraph_text is None and insert_paragraph_at is None:
        if paragraph_index < 0 or paragraph_index >= para_count:
            return f"[ERROR] 段落索引 {paragraph_index} 超出范围 (0-{para_count - 1})"
        p = paragraphs[paragraph_index]._element
        p.getparent().remove(p)
        doc.save(file_path)
        return f"[OK] 已删除 DOCX 段落 {paragraph_index}: {file_path}"

    # 替换段落
    if paragraph_index is not None and new_paragraph_text is not None:
        if paragraph_index < 0 or paragraph_index >= para_count:
            return f"[ERROR] 段落索引 {paragraph_index} 超出范围 (0-{para_count - 1})"
        paragraphs[paragraph_index].text = new_paragraph_text
        doc.save(file_path)
        return f"[OK] 已替换 DOCX 段落 {paragraph_index}: {file_path}"

    # 插入段落
    if insert_paragraph_at is not None and insert_text is not None:
        if insert_paragraph_at < 0:
            insert_paragraph_at = 0
        if insert_paragraph_at > para_count:
            insert_paragraph_at = para_count
        if insert_paragraph_at == para_count:
            doc.add_paragraph(insert_text)
        else:
            ref_para = paragraphs[insert_paragraph_at]
            new_para = ref_para.insert_paragraph_before(insert_text)
        doc.save(file_path)
        return f"[OK] 已在 DOCX 段落 {insert_paragraph_at} 前插入新段落: {file_path}"

    return "[ERROR] docx update 操作缺少必要参数"


# ---------- pptx ----------
def _read_pptx(file_path: str, max_chars: int | None = None, from_end: bool = False,
               max_images: int | None = None) -> tuple[str, list[bytes]]:
    text, images = Documents_process.process_pptx(file_path)
    if max_images is not None and max_images >= 0 and len(images) > max_images:
        images = images[:max_images] if not from_end else images[-max_images:]
    img_note = f"\n(含 {len(images)} 张内嵌图片)" if images else ""
    result = f"=== {file_path} ===\n{text}{img_note}"
    result = _apply_text_truncation(result, max_chars, from_end)
    return result, images


# ---------- pdf ----------
def _read_pdf(file_path: str, max_chars: int | None = None, from_end: bool = False,
              max_images: int | None = None) -> tuple[str, list[bytes]]:
    text, images = Documents_process.process_pdf(file_path)
    # 截取图片
    if max_images is not None and max_images >= 0 and len(images) > max_images:
        images = images[:max_images] if not from_end else images[-max_images:]
    img_note = f"\n(含 {len(images)} 张内嵌图片)" if images else ""
    result = f"=== {file_path} ===\n{text}{img_note}"
    result = _apply_text_truncation(result, max_chars, from_end)
    return result, images


# ---------- image ----------
def _read_image(file_path: str, max_images: int | None = None,
                from_end: bool = False) -> tuple[str, list[bytes]]:
    """读取独立图片文件，将图片字节放入 artifact 返回。

    输入:
        file_path: 图片文件路径。
        max_images: 最多返回的图片数（对单张图片通常为1或None）。
        from_end: 截取方向。

    输出:
        (文本描述, 图片字节列表)。
    """
    _, images = Documents_process.process_image(file_path)
    if max_images is not None and max_images >= 0 and len(images) > max_images:
        images = images[:max_images] if not from_end else images[-max_images:]
    img_note = f"(含 {len(images)} 张图片)" if images else ""
    result = f"=== {file_path} ===\n[图片文件]{img_note}"
    return result, images


# ---------- audio ----------
def _read_audio(file_path: str) -> tuple[str, list[bytes]]:
    """读取音频文件，通过 ASR 转为文字返回。

    输入:
        file_path: 音频文件路径。

    输出:
        (ASR识别文本, 空列表)。
    """
    try:
        asr_text = audio_file_to_text(file_path)
    except Exception as e:
        return f"[ERROR] ASR 识别失败: {str(e)}", []

    result = f"=== {file_path} ===\n[音频文件 ASR 识别结果]\n{asr_text}"
    return result, []


# ---------- Dispatch ----------
def _dispatch_read(doc_type: str, file_path: str, **kwargs) -> tuple[str, list[bytes]]:
    max_chars = kwargs.get("max_chars")
    from_end = kwargs.get("from_end", False)
    max_images = kwargs.get("max_images")
    grep_pattern = kwargs.get("grep_pattern")

    if doc_type == "plain_text":
        return _read_plain_text(file_path, grep_pattern=grep_pattern, max_chars=max_chars, from_end=from_end), []
    elif doc_type == "tabular":
        return _read_tabular(file_path, kwargs.get("sheet_name"), max_chars=max_chars, from_end=from_end), []
    elif doc_type == "docx":
        return _read_docx(file_path, max_chars=max_chars, from_end=from_end, max_images=max_images)
    elif doc_type == "pptx":
        return _read_pptx(file_path, max_chars=max_chars, from_end=from_end, max_images=max_images)
    elif doc_type == "pdf":
        return _read_pdf(file_path, max_chars=max_chars, from_end=from_end, max_images=max_images)
    elif doc_type == "image":
        return _read_image(file_path, max_images=max_images, from_end=from_end)
    elif doc_type == "audio":
        return _read_audio(file_path)
    return f"[ERROR] read 不支持文档类型: {doc_type}", []


def _dispatch_create(doc_type: str, file_path: str, content: str, **kwargs) -> str:
    if doc_type == "plain_text":
        return _create_plain_text(file_path, content)
    elif doc_type == "tabular":
        return _create_tabular(file_path, content)
    elif doc_type == "docx":
        return _create_docx(file_path, content, kwargs.get("image_paths"))
    return f"[ERROR] create 不支持文档类型: {doc_type} (仅支持 plain_text、tabular、docx)"


def _dispatch_update(doc_type: str, file_path: str, **kwargs) -> str:
    if doc_type == "plain_text":
        return _update_plain_text(
            file_path,
            kwargs.get("line_number"),
            kwargs.get("new_line_text"),
            kwargs.get("insert_before_line"),
            kwargs.get("insert_text"),
        )
    elif doc_type == "tabular":
        return _update_tabular(
            file_path,
            kwargs.get("row"),
            kwargs.get("col"),
            kwargs.get("cell_value"),
            kwargs.get("sheet_name"),
        )
    elif doc_type == "docx":
        return _update_docx(
            file_path,
            kwargs.get("paragraph_index"),
            kwargs.get("new_paragraph_text"),
            kwargs.get("insert_paragraph_at"),
            kwargs.get("insert_text"),
            kwargs.get("table_index"),
            kwargs.get("docx_table_row"),
            kwargs.get("docx_table_col"),
            kwargs.get("docx_table_cell_value"),
            kwargs.get("insert_image_path"),
        )
    return f"[ERROR] update 不支持文档类型: {doc_type} (仅支持 plain_text、tabular、docx)"


class DocTool(BaseTool):
    args_schema: type[BaseModel] = DocToolInput
    name: str = "doc_tool"
    description: str = (
        "文档操作工具。支持纯文本(plain_text)、表格(tabular: CSV/XLSX/XLS)、Word(docx)、PowerPoint(pptx)、PDF 文档的增删改查，"
        "以及图片(image)、音频(audio)的读取。\n\n"
        "【支持的操作】\n"
        "- read: 读取文档内容并带行号返回。支持所有类型。\n"
        "  * 内容控制: max_chars(最大字符数)、from_end(True从末尾截取)、max_images(docx/pdf/pptx 最大图片数)。\n"
        "  * plain_text grep: grep_pattern 为正则表达式，仅返回匹配的行(含原始行号)。\n"
        "- create: 创建新文档。仅支持 plain_text、tabular、docx。\n"
        "  * plain_text/docx: content 为要写入的文本。\n"
        "  * docx 可附带图片: image_paths 为图片文件路径列表，图片插入到文本段落后。\n"
        "  * tabular(.csv): content 为 CSV 格式文本。\n"
        "  * tabular(.xlsx): content 为 JSON 二维数组字符串，如 '[['A','B'],['1','2']]'。\n"
        "- update: 修改文档。仅支持 plain_text、tabular、docx。\n"
        "  * plain_text 行级操作: line_number + new_line_text 替换行；\n"
        "    line_number 单独使用删除行；\n"
        "    insert_before_line + insert_text 插入行。\n"
        "  * tabular 单元格操作: row + col + cell_value 更新单元格。(1-indexed，1=表头)\n"
        "  * docx 段落操作: paragraph_index + new_paragraph_text 替换段落；\n"
        "    paragraph_index 单独使用删除段落；\n"
        "    insert_paragraph_at + insert_text 插入段落。\n"
        "  * docx 表格操作: table_index + docx_table_row + docx_table_col + docx_table_cell_value。\n"
        "  * docx 插入图片: insert_paragraph_at + insert_image_path 在指定段落前插入图片。\n"
        "- delete: 删除文档。支持所有类型。\n\n"
        "【文档类型识别】\n"
        "  plain_text: .txt .json .md .cpp .c .py .m .java .html .svg 等\n"
        "  docx: .docx\n"
        "  pptx: .pptx（read 提取文本和图片；delete 删除）。"
        "  pdf: .pdf\n"
        "  tabular: .csv .xlsx .xls\n"
        "  image: .png .jpg .jpeg .bmp .gif .webp .ico\n"
        "  audio: .wav .mp3 .flac .ogg .m4a .aac 等（仅 read，自动 ASR 转文字）\n\n"
        "【重要：文件路径】\n"
        "  创建新文件时，默认存放于系统提示词中指定的工作空间目录（WORKSPACE_DIR）。\n"
        "  请始终使用完整的绝对路径，如 'C:\\\\Users\\\\...\\\\workspace\\\\my_script.py'。\n"
    )
    response_format: Literal["content", "content_and_artifact"] = "content_and_artifact"

    def _run(self, file_path: str, action: str, **kwargs) -> tuple[str, dict]:
        try:
            doc_type = _classify(file_path) if action != "read" or os.path.isfile(file_path) else _classify_by_ext(file_path)
            if not doc_type:
                return (f"[ERROR] 无法识别文档类型: {file_path}", {})

            if action == "read":
                if not os.path.isfile(file_path):
                    return (f"[ERROR] 文件不存在: {file_path}", {})
                text, images = _dispatch_read(doc_type, file_path, **kwargs)
                # 将图片放入 artifact
                artifact = {}
                if images:
                    import base64
                    from agent.utils.image_utils import image_bytes_to_openai_image_url_part
                    artifact["images"] = images
                    artifact["file_path"] = file_path
                    # 为前端展示生成 data URL 列表
                    artifact["image_urls"] = [
                        f"data:image/png;base64,{base64.b64encode(img).decode('ascii')}"
                        for img in images
                    ]
                return (text, artifact)

            elif action == "create":
                content = kwargs.get("content") or ""
                if not content:
                    return ("[ERROR] create 操作需要提供 content 参数", {})
                # content 已作为位置参数传递，从 kwargs 中移除避免重复
                extra = {k: v for k, v in kwargs.items() if k != "content"}
                return (_dispatch_create(doc_type, file_path, content, **extra), {})

            elif action == "update":
                return (_dispatch_update(doc_type, file_path, **kwargs), {})

            elif action == "delete":
                return (_delete_file(file_path), {})

            return (f"[ERROR] 未知操作: {action}", {})
        except Exception as e:
            return (f"[ERROR] doc_tool 执行 {action} 操作时出错: {str(e)}", {})

    async def _arun(self, **kwargs) -> tuple[str, dict]:
        import asyncio
        return await asyncio.to_thread(self._run, **kwargs)
